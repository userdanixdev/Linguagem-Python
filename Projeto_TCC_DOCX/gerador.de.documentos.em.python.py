"""
gera_sgbd_docx.py
Gera um documento Word com o trabalho:
"Um levantamento técnico sobre Sistemas Gerenciadores de Banco de Dados"

Como usar:
1. Instale a biblioteca python-docx:
   pip install python-docx

2. Execute:
   python gera_sgbd_docx.py

O arquivo será salvo no mesmo diretório com o nome:
Um_levantamento_tecnico_sobre_SGBD_Daniel_Martins_Franca.docx
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import ast
import inspect

def cria_capa(doc, titulo, autor, curso, instituicao, local_data):
    # inserir espaçamento superior na capa
    p_top = doc.add_paragraph()
    p_top.add_run("\n\n\n\n\n\n\n")
    # título centralizado e em negrito
    p_titulo = doc.add_paragraph()
    run = p_titulo.add_run(titulo)
    run.bold = True
    run.font.size = Pt(18)
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # espaçamento
    doc.add_paragraph("\n\n\n\n")

    p_autor = doc.add_paragraph(autor)
    p_autor.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_curso = doc.add_paragraph(curso)
    p_curso.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_instituicao = doc.add_paragraph(instituicao)
    p_instituicao.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_local = doc.add_paragraph(local_data)
    p_local.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

def adiciona_rodape(doc, nome, ra):
    secao = doc.sections[0]
    footer = secao.footer
    # Garantir que o rodapé fique realmente embaixo
    secao.footer_distance = Pt(35)

    rodape = secao.footer
    paragrafo = footer.paragraphs[0]
    paragrafo.clear()

    # ========== LINHA HORIZONTAL ==========
    p_pr = paragrafo._p.get_or_add_pPr()
    bordas = OxmlElement('w:pBdr')
    linha = OxmlElement('w:bottom')
    linha.set(qn('w:val'), 'single')
    linha.set(qn('w:color'), '000000')
    linha.set(qn('w:sz'), '6')  # espessura
    bordas.append(linha)
    p_pr.append(bordas)
    # Sobe ainda mais o texto do rodapé
    espaco = OxmlElement("w:spacing")
    espaco.set(qn("w:before"), "100")   # <--- SUBINDO MAIS
    p_pr.append(espaco)
    # ================================
    #  COLUNA ESQUERDA → nome + RA
    # ================================
    texto_esq = paragrafo.add_run(f"{nome} | RA: {ra}")

    # ================================
    #  TABULAÇÃO PARA EMPURRAR À DIREITA
    # ================================
    paragrafo.add_run("\t\t")

    # ================================
    #  CAMPO DE NÚMERO DA PÁGINA
    # ================================
    campo_pag = OxmlElement('w:fldSimple')
    campo_pag.set(qn('w:instr'), "PAGE")  # campo real
    paragrafo._p.append(campo_pag)

    
    # Configurar tabulação à direita para alinhar
    ppr = paragrafo._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "10500")  # posição na margem direita
    tabs.append(tab)
    ppr.append(tabs)

    # ================================
    #  OCULTAR NA PRIMEIRA PÁGINA
    # ================================
    secao.different_first_page_header_footer = True
    rodape_primeira = secao.first_page_footer
    rodape_primeira.paragraphs[0].text = ""

def aplica_fonte_padrao(doc, nome_fonte="Times New Roman", tamanho=12):
    """
    Define a fonte padrão do documento (para parágrafos novos).
    Nota: python-docx não altera facilmente estilos existentes em todos os casos,
    mas podemos ajustar os estilos principais.
    """
    style = doc.styles['Normal']
    style.font.name = nome_fonte
    # necessário para compatibilidade com alguns sistemas:
    r = style._element.rPr.rFonts
    r.set(qn('w:eastAsia'), nome_fonte)
    style.font.size = Pt(tamanho)

def adiciona_secao_titulo_e_paragrafo(doc, titulo, texto):
    doc.add_heading(titulo, level=1)
    parag = doc.add_paragraph(texto)
    parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def adiciona_sumario(doc, main):

    """
    Gera um sumário com base nos headings encontrados dentro da função main passada como objeto.
    
    Parâmetros:
        doc: objeto Document do python-docx
        func_main: função main (ex: main)
    """
    # Obter o código-fonte da função main em string
    codigo = inspect.getsource(main)
    
    # Parse do código com AST
    tree = ast.parse(codigo)
    headings = []

    # Percorrer nós da AST
    for node in ast.walk(tree):
        # doc.add_heading("Título", level=...)
        if isinstance(node, ast.Call):
            if isinstance(node.func, ast.Attribute):
                if node.func.attr == "add_heading" and len(node.args) >= 1:
                    if isinstance(node.args[0], ast.Constant):
                        headings.append(node.args[0].value)
            # adiciona_secao_titulo_e_paragrafo(doc, "Título", ...)
            elif isinstance(node.func, ast.Name) and node.func.id == "adiciona_secao_titulo_e_paragrafo":
                if len(node.args) >= 2 and isinstance(node.args[1], ast.Constant):
                    headings.append(node.args[1].value)

    if not headings:
        return
    
        # --- Inserir sumário sem folha extra ---
    # Colocar título centralizado
    titulo = doc.add_paragraph("SUMÁRIO")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_titulo = titulo.runs[0]
    run_titulo.bold = True
    run_titulo.font.size = Pt(12)
    doc.add_paragraph("")  # pequeno espaçamento

    for texto in headings:
        p = doc.add_paragraph()
        p.style = "Normal"

        # Tabulação com líder de pontos
        ppr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        ppr.append(tabs)

       # Adicionar texto e espaço reservado para número da página
        run = p.add_run(texto)
        run.font.size = Pt(12)
        p.add_run("\t")  # tabulação para líder de pontos
        p.add_run(" ")   # espaço para número de página manual

    doc.add_page_break()
             


def main():
    titulo = "Noções básicas de Geodésia"
    autor = "Daniel Martins França"
    curso = "Curso de Extensão"
    instituicao = "Faculdade Metropolitana"
    local_data = "Brasília – DF, novembro de 2025"
    nome_arquivo = "geografia.docx"
    doc = Document()
    
    
    aplica_fonte_padrao(doc, nome_fonte="Times New Roman", tamanho=12)
    # Capa
    cria_capa(doc, titulo, autor, curso, instituicao, local_data)
    adiciona_sumario(doc, main)
    
    adiciona_rodape(doc,nome="Daniel Martins França", ra="2101562")
    doc.save(nome_arquivo)
    print(f"Documento '{nome_arquivo}' criado com rodapé criado com sucesso.")

    # Introdução
    texto_intro = (
"""Diversos modelos foram adotados ao longo da história. Não existe modelo errado, tudo depende do contexto histórico e da aplicação.
	Assim, o primeiro modelo é chamado de geoidal, é o mais aproximado da forma real, podendo ser determinado pelas medidas gravimétricas,
ou seja, medidas da força da gravidade, explicado a seguir. O geoide não pode ser definido matematicamente pois é afetado pelas variações da densidade dos
elementos constituintes da crosta terrestre. Além da distribuição irregular das massas terrestres e oceânicas.
	Devido às irregularidades da superfície terrestre, utilizam-se modelos para a sua representação, mais simples, outros regulares e geométricos e que se
aproximam da forma real para efetuar os cálculos. Cada modelo tem a sua aplicação e quanto mais complexa a figura empregada para a representação da Terra,
mais complexo serão os cálculos sobre esta superfície. A forma da Terra gira em torno do seu eixo e movendo-se dentro do sistema solar. E o resultado da interação
de forças internas e externas tais como: gravidade, força centrífuga e a constituição diferente dos materiais que a formam a geoide ao longo de milhares de anos.
As forças tectônicas provocam modificações na superfície, que se traduzem por irregularidades topográficas, sobre as quais são realizados os mapeamentos, 
medições e estudos das mais variedades.
"""
    )
    adiciona_secao_titulo_e_paragrafo(doc, "Estudos Geodésicos:", texto_intro)

    # Desenvolvimento (vários parágrafos)
    doc.add_heading("Geoide:", level=1)

    p1 = (
    """É a forma física real, que sofre frequentes alterações devido a natureza (campo gravitacional do planeta terra, movimentos tectônicos, condições climáticas,
erosões, etc) e à ação do homem, portanto, não serve para definir forma sistemática da Terra. Portanto, a zona de contato da superfície terrestre topográfica e
o geoide que defini-se o nível zero das altitudes. Em resumo, a força gravitacional que age sobre a terra irá definir a forma irregular do planeta. O equilíbrio
desse potencial gravitacional é o que gera a forma física da Terra. Mantendo os oceanos e a divisão dos continentes com uma relação de equilíbrio do planeta.
	Portanto, a geoide é a superfície equipotencial gravitacional que mais se aproxima da superfície formada pelo prolongamento dos oceanos (nível médio dos mares)
sob os continentes. Essa supefície sofre variações conforme ocorrem alterações no campo gravitacional terrestre e, portanto, não segue leis matemáticas que permitam
um modelamento da Terra. Ainda assim é empregado como referência para a determinação das altitudes. """

    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p2 = (
        "A Gravimetria é um método de medida da gravidade utilizado em diversos pontos distribuidos do planeta Terra. Para que apartir dessas medições se alcançem um modelo"
"de ação gravitacional equilibrado. Portanto, a Gravimetria é um método da Geodésia física que determina os níveis do campo gravitacional da Terra e, com isto,"
"determinar o geoide. A densidade de pontos é muito importante para a determinação do geoide. Quanto mais pontos de medição gravimétrica existirem na superfície" 
"terrestre, mais precisa é a figura geoidal. A incidência nos pontos gravitacionais é perpendicular a superfície da terra deverá ser ortométrica (distância contada sobre a vertical)"
"As altitudes ortométricas (ângulo reto) a ação da gravidade nesses pontos em que teremos a determinação de uma superficíe equipontencial. Essa superfície"
"equipotencial é chamada de um modelo geoide. Em média, coincide com o valor médio do nível médio das águas do mar, por isso é usado para medições"
"de altitudes  (altimetria). A superfície geoidal é mais irregular que qualquer outra superfície. A superfície varia entre os +8.850 m (Monte Everest) e -11.000m "
"(Fossa das Marianas). O geoide varia apenas cerca de +-100 m além das superfície do elipsoide de referência."

    )
    doc.add_paragraph(p2).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("Elipsoidal:", level=1)

    p1 = (
        "Como o geoide é uma superfície matematicamente indefinida, as reduções ou transferências de dados a ele são inconsistentes, e para um mapeamento preciso de "
"grandes áreas é necessária a consideração de uma figura geométrica regular. Assim, chega-se na figura matematicamente definida como um elipsoide."
"A elipsoide é a representação matemática mais próximas da forma da Terra. Esse método é utilizado para GNSS e Geodésia Celeste, inclusive. (tecnologias novas)"
"O modelo Elipsoide de revolução é o nome que se dá a uma figura gerada pela rotação de uma elipse sobre um de seus eixos(de norte a sul da esfera). Eixo polar (Norte e Sul) "
"é menor do que o eixo equatorial ( horizontal ) que por sua vez também possui semi-eixo. A relação entre o semi-eixo maior (equatorial) e o semi-eixo menor (polar)"
"é o que define a chamada excêntricidade em que ocorre o achatamento da elipsoide. A posição deste elipsoide em relação à Terra, bem como sua forma e tamanho,"
"constituem um conjunto de parâmetros definidos a partir do elipsoide ao seu ajustamento com a geoide onde são usualmente denominados Datum Geodésico. "
"Uma elipsoide de revolução pode ser determinado para melhor se adaptar a uma região, país ou continente, evitando a ocorrência de desníveis geoidais muito exagerados."
"Em geral, cada país adotou um elipsoide como referência para os trabalhos geodésicos e topográficos. Atualmente existem pelo menos 50 elipsoides existentes"
"utilizados pelo mundo, sendo alguns mais conhecidos: SAD-69, WGS-84, SIRGA-2000. "
"Essas siglas, sempre encontradas no mapas, são elipsoides de revolução. Elas ilustram a referência utilizada para usar em trabalhos cartográficos."
"Sendo assim, a superfície elipsoidal, que é obtido a partir do elipsoide de revolução, é uma figura matemática constante. Então, cada país ou região, para estudos"
"de mapeamento, por exemplo, é estabelecer qual é o elipsoide de referência para o mapeamento naquele local."
"	O semi-eixo do WGS 1984 possui um eixo equatorial, em metros, 6.378.137. O eixo polar possui, em metros, 6.356.752.3142."
"Foram usados equipamentos novos como GPS, GNSS, RTK para o datum WGS 1984."

    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("DATUM:", level=1)
    p1 = (
        "Datum: modelo matemático teórico da representação da superfície da Terra ao nível do mar. O plural de DATUM é data."

"Modelos empregados em projetos:"

"	Córrego Alegre "
"	WGS84 <- Geocêntrico ( O ponto de origem coincide com o Centro de Massa da Terra, validade GLOBAL )"
	"SAD69"
	"SIRGAS2000 <- Oficial do Brasil ( Geocêntrico , o ponto de origem coincide com o Centro de Massa da Terra , validade GLOBAL )"
	"	O SIRGAS2000 foi constituído para representar melhor a América do Sul."
"O que são códigos EPSG?"

"	Os códigos foram feitos pelos europeus: European Petroleum Survey Group (ESPG) Coleção de sistemas de coordenadas existentes e atribuiram um código "
"numérico para identifica. Ele, por exemplo, identifica o DATUM + Sistema de coordenada."
"Não existe sistema de coordenadas geográficas UTM, o sistema de coordenada plana que é UTM. O sistem de coordenada geográfica é outro sistema."

"DATUM + Sistema de coordenadas"

"O que é DATUM: Modelo matemático teórico da representação da superfície da Tera ao nível do mar. Existem diferentes tipos de modelos DATUM."
"O oficial usado no Brasil é o SIRGAS 2000. ( O IBGE oficializou a partir de 2015, o datum SIRGAS2000)"
 "	Todo produto cartográfico gerado aqui no Brasil, deve utilizar como DATUM, o SIRGAS2000."
"Realmente, o WGS84 e o SIRGAS2000 são modelos DATA geocêntricos. O SAD69 é topocêntrico. Na legislação brasileira, o data oficial é o SIRGAS2000."
"Eles são semelhantes, com poucas diferenças principalmente em aplicações de mapeamento. A referência do SIRGAS2000 é mais adequado para região do Brasil. "

	"E o sistema de coordenadas?"

"No âmbito de geração de mapas e geoprocessmento existem basicamente dois tipos principais de coordenadas. A coordenada UTM ou a coordenada geográfica ou geodésica."
"As coordenadas UTM são expressas em unidades métricas ( em metros ). Já as coordenadas geográficas/geodésica são expressas em grau decimal ou grau minuto/segundo."

"	Caso utilize o mesmo sistema de coordenada, por exemplo, o DATUM utilizado for diferente, a localização do ponto será em um local diferente/distinto."
"Cada DATUM tem um modelo matemático diferente. Por isso, a localização será alterada. Sendo assim, ao fazer um trabalho de mapeamento deverá utilizar o DATUM"
"e o sistema de coordenadas adequado e compatível no projeto para não haver problemas posteriores."

"Por exemplo, no Brasil temos as divisões estadual. Nela existem linhas verticais, chamadas meridianos. Sendo assim, está sendo utilizado sistema de coordenada UTM"
"Cada meridiano são zonas ou fusos. Se a área a ser mapeada está dentro de um único fuso, pode optar por utilizar coordenadas UTM ou geográficas. Mas se a área"
"de estudo pegar uma área maior que o da zona, não deverá utilizar coordenadas UTM e sim, coordenadas geográficas. No Brasil, 4 estados estão inseridos em um único"
"fuso. Portanto, podemos utilizar coordenadas métricas UTM. Ex: Ceará (zona 24 SUL), Sergige (zona 24 SUL), Espírito Santo (zona 24 Sul) e Santa Catarina (zona 22 SUL)."
 "	A coordenadas UTM, os cálculos métricos são mais rápidos. O estado do AMAZONAS, várias fusos, várias zonas. Sendo assim, em geral, por padrão, os programas"
"de geoprocessamento irão calcular distâncias e áreas que vão levar em consideração: o sistema de coordenadas."
	"Agora se for analisar um município dento de um estado, podemos mapear por coordenadas UTM."

        "A posição do elipsoide em relação a terra, bem como sua forma e tamanho, constituem um conjunto de parâmetros que são usualmente denominados de DATUM"
"GEODÉSICOS. Datum corresponde a um ponto ou plano de referência para levantamento verticais e horizontais, os quais estabelecem as posições de feições sobre"
"a terra."

    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("REDES de sistema geodésico de referência ALTIMÉTRICO E PLANIMÉTRICO::", level=1)
    doc.add_heading("DATUM PLANIMÉTRICO", level=1)


    p1 = (
        "Nos referenciais clássicos a determinação do ponto de origem (topocêntrico) se da pelo relacionamento entre coordenadas geodésicas e astronômicas"
"utilizando as condições de La Place: "

"		     Ψa - Ψg = e (componente meridiana)"
		"(λa-λg).cos(g) = η (componente 1° vertical)	"

"No ponto de origem do sistema eram realizadas medidas astronômicas para determinar latitude, longitude e azimute astronômico o qual se "
"determinava as coordenadas geodésicas de partidade para o referencial."
        
        "É o ponto em uma região de melhor coincidência do elipsoide de referência ao geoide, onde o desvio da vertical é nulo ou mínimo. O conhecimento do"
"desvio da vertical é importante para a escolha do datum planimétrico do sistema geodésico de apoio ao levantamento cartográfico de um país."
"O datum planimétrico ou horizontal é definido como a forma e tamanho de um elipsoide, bem como sua posição relativa ao geoide. Cada país adota um elipsoide"
"como referência, o qual se aproxima do geoide da região considerada. O Brasil adota o elipsoide do Sistema Geodésico de Referência de 1980. (GRS80)"
"A adoção de um referencial topocêntrico, onde o ponto de origem e orientação está na superfície terrestre e não no centro da terra."
"A rede planimétrica clássica, onde possui também, milhares de redes, são compostos por vértices de triangulação em que devem ser intervisíveis. "
"Usado muito até os anos 90 até implantarem as estações GNSS. Hoje em dia, o Brasil também possui uma rede de monitoramento contínuo dos sistemas GNSS em vários"
"pontos do Brasil."

"	O sistema geodésico brasileiro era planimétrica clássica, o SAD 69 (South American Datum 1969), o modelo da Terra era o elipsóide internacional de 1967 (GRS 67)"
"Sua origem era topocêntrica com vértice de Chuá. "

"Hoje o sistema de referência é geocêntrico para as Américas, SIRGA2000, usado o modelo geoide, um elipsoide "
"internacional de 1980 (GRS 80), a sua origem era o centro da massa da Terra (Elipsoide Geocêntrico)"
"Com a evolução nas tecnologias, houve reajuste no SAD69(1996). Com o advento do SIRGAR 2000 a definição/orientação do SIRGAS2000 é geocêntrica: adota um referencial"
"que tem a origem dos seus três eixos cartesianos localizada no centro de massa da Terra. Parecido com o WGS84, é internacional.  Adota o mesmo sistema geodésico de "
"referência, geocêntrico também, e tem o objetivo de fornecer posicionamento e navegação em qualquer parte do mundo. O Google Earth utilizada o datum WGS84."

    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("SISTEMA GEODÉSICO DE REFERÊNCIA:\nREDES de sistema geodésico de referência ALTIMÉTRICO E PLANIMÉTRICO:", level=1)
    
    p1 = (
    "SGR são estruturas constituídas por estações geodédicas materializadas na superfície terrestre, com coordenadas que servem de referência para os"
"levantamentos topográficos e geodésicos a serem realizados no território de interesse. Isso, constitui-se a infra-estrutura de referência para novos"
"posicionamentos a serem efetuados. No Brasil, o órgão responsável em estabelecer esses sistemas de referência é o IBGE.    "
"O sistema geodésico de referência brasileira possui milhares de redes altimétricas ou verticais. As 'RRNN' são as referências de nível, onde os pontos cuja a altitude "
"ao nível do mar é conhecida. Para uso de mapas topográficos, é utilizado o método ortométrico, em que calcula a distância entre o geoide até a superfície terrestre,"
"denominada 'H'. Além disso, temos a altitude elipsoidal 'h', sendo a altitude da elipsoide de revolução até a superfície terrestre, as medições são feitas com GPS."
"E o 'N' é a ondulação geoidal, entre o geoide e o elipsoide. Sendo assim, temos a fórmula para obter a altitude ortométrica: 'H ≅ h + N'"
"O datum altimétrico usado no Brasil oficial é o Marégrafo de Imbituba/SC, desde 1958."
"O datum planimétrico ou horizontal é definido como a forma e tamanho de um elipsoide, bem como sua posição relativa ao geoide. Cada país adota um elipsoide"
"como referência, o qual se aproxima do geoide da região considerada. O Brasil adota o elipsoide do Sistema Geodésico de Referência de 1980. (GRS80)"
"A adoção de um referencial topocêntrico, onde o ponto de origem e orientação está na superfície terrestre e não no centro da terra."
"om a evolução nas tecnologias, houve reajuste no SAD69(1996). Com o advento do SIRGAR 2000 a definição/orientação do SIRGAS2000 é geocêntrica: adota um referencial"
"que tem a origem dos seus três eixos cartesianos localizada no centro de massa da Terra. Parecido com o WGS84, é internacional.  Adota o mesmo sistema geodésico de "
"referência, geocêntrico também, e tem o objetivo de fornecer posicionamento e navegação em qualquer parte do mundo. O Google Earth utilizada o datum WGS84."

"	A rede planimétrica clássica, onde possui também, milhares de redes, são compostos por vértices de triangulação em que devem ser intervisíveis. "
"Usado muito até os anos 90 até implantarem as estações GNSS. Hoje em dia, o Brasil também possui uma rede de monitoramento contínuo dos sistemas GNSS em vários"
"pontos do Brasil."

	"Os sistema geodésico brasileiro era planimétrica clássica, o SAD 69 (South American Datum 1969), o modelo da Terra era o elipsóide internacional de 1967 (GRS 67)"
"Sua origem era topocêntrica com vértice de Chuá. "

"Hoje o sistema de referência é geocêntrico para as Américas, SIRGA2000, usado o modelo geoide, um elipsoide "
"internacional de 1980 (GRS 80), a sua origem era o centro da massa da Terra (Elipsoide Geocêntrico)"

"As três superfícies da Geodésia:"

"A superfície verdadeira (superfície equipotencial de referência): geoide (datum vertical);"
"A superfície matemática: elipsoide (datum horizontal e referência vertical); "
"A superfície física: terrestre (onde são realizadas as medições). "
    #imagem
"Onde temos :"

"'h': altitude geométrica (elipsoidal) cuja distância vertical medida sobre a normal (perpendicular ao elipsoide) entre o ponto medido e uma superfície de"
"referência elipsoidal."
"'H': altitude ortométrica (geoidal) cuja distância é medida da geoide ou nível do mar até a superfície física, independe do elipsoide de referência e tem "
"seu maior significado físico."
"'N': a ondulação gravitacional geoidal cuja distância entre a superfície geoidal (verdadeira) à superfície elipsoidal de referência, perpendicular sobre"
"um segmento de reta entre os pontos."

 )

    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("DATUM ALTIMÉTRICO", level=1)

    p1 = (
       " Esse método utiliza-se uma superfície de referência e uma superfície de nível chama-se de datum vertical ou altimétrico."
       "e refere-se as altitudes. A determinação do datum vertical envolve um marégrafo ou uma rede de marégrafos para determinar o nível médio do mar (NMM) em um ponto específico da costa.         "
        "Ponto fixo fundamentado e solidamente materializado, cuja altitude sobre o nível do mar é utilizado como partida de referência das altitudes que"
"determinam os nivelamentos."
    
 )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("SISTEMAS GEODÉSICOS REGIONAIS E GLOBAIS", level=1)
    
    
    p1 = ("Os sistemas geodésicos regionais é um referencial adaptado a uma região (país ou continente) devido à limitação dos métodos de posicionamento utilizado."
"Permite a possibilidade de existência de mais de um sistema de referência em cada região ou país. Por exemplo: Chuá-Astro Datum, South American Datum 1969 (SAD69)"
"e Córrego Alegre."
"O advento dos satélites artificiais para posicionamento possibilitou o desenvolvimento prático dos sistemas de referência geocêntricos, como "
"por exemplo, o WGS84 e o ITRS (International Terrestrial Reference System) em suas mais diversas realizações e densificações."
"	O ITRS2000 é uma densificação e deu origem ao sistema SIRGAS2000, sendo ambos compatíveis entre os sistemas. Possuem as mesmas características."
	"Os sistemas geodésico global são adequados às modernas técnicas de posicionamento, possibilitando levantamentos globais. Como exemplo pode-se destacar"
"os sistemas globais de navegação por Satélite (GNSS - Global Navigation Satellite System). A origem do sistema é o centro de massa da Terra ( geocêntrico )"
"Exemplos são: World Geodetic System 1984 - WGS84; Internacional Terrestrial Reference System - ITRS e o Sistema de Referência Geocêntrico para as Américas"
"-SIRGAS 2000."

    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading(" SISTEMA GEODÉSICO GLOBAL: International Terrestrial Reference System (ITRS):", level=1)
    
    
    p1 = (" ITRS2000 é uma densificação e deu origem ao sistema SIRGAS2000, sendo ambos compatíveis entre os sistemas. Possuem as mesmas características."
	"Os sistemas geodésico global são adequados às modernas técnicas de posicionamento, possibilitando levantamentos globais. Como exemplo pode-se destacar"
"os sistemas globais de navegação por Satélite (GNSS - Global Navigation Satellite System). A origem do sistema é o centro de massa da Terra ( geocêntrico )"
"Exemplos são: World Geodetic System 1984 - WGS84; Internacional Terrestrial Reference System - ITRS e o Sistema de Referência Geocêntrico para as Américas"
"-SIRGAS 2000."

"International Terrestrial Reference System (ITRS):"
	
"	O ITRS é um sistema de referência moderno que define matematicamente como representar posições na Terra."
"Ele rotaciona com a Terra em seu movimento no espaço. Cada realização do ITRS é especificada em coordenadas cartesianas X,Y e Z."

"	- Eixo Z aponta na direção do CTP( Conventional Terristrial Pole = Polo Norte Terrestre Convencional)"
"	- Eixo X aponta para a direção média do meridiano de Greenwich"
	"- Eixo Y aponta para o leste, perpendicular ao eixo X "

"O ITRS é realizado a partir de um conjunto de coordenadas e velocidades observadas por GNSS,LLR, GPS, SLR, VLBI e DORIS. "
"Cada realização é denominada de International Terrestrial Reference Frame (ITRF), na prática + ano de realização."
"É atualizado periodicamente (2000,2008,2014,2020...)"
"Em resumo, os posicionamentos GNSS modernos estão, na prática, referidos ao ITRS/ITRF."

"	Os ITRS são importantes para mapas globais, referência de posicionamento preciso e controle de placas tectônicas."
"Eles monitoram a elevação do nível do mar, mudanças climáticas e deformações na terra."
    )

    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("SISTEMA GEODÉSICO BRASILEIRO:", level=1)
       
    
    p1 = (
"""O sistema Geodésico Brasileiro adotou oficialmente 3 referênciais:
Córrego Alegre; SAD69; SIRGA2000
Sistemas de referência geodésico usados no Brasil:
Córrego Alegre (1911): Ajustamento da rede planimétrica na década de 40
- Época de Referência: 1911- Elipsoide: Hayford (Internacional 1924)
- Semi-eixo maior: a = 6.378.388 m- Achatamento: f = 1/297
- Origem: Ajustado a observações astronômicas locais no Brasil ( topocêntrico )
- Ondulação geoidal: N = 0- Observação: Foi o primeiro sistema oficial, usado até os anos 1970.
- Coordenadas do ponto de origem:
	A = 19"50'15,14"s e B = 48"57'42,75"W

Os componentes de desvio da vertical foram nulas, dessa forma a normal e a vertical são coincidentes nesse ponto, sendo assim, o elipsoide
é paralelo ao geoide no vértice Córrego Alegre. A ondulação geoidal também é nula, fazendo com que o elipsoide seja coindidente com o geoide
no ponto origem. As realizações do CA foram nos anos de 1961,1970 e 1972. Para fins de transformações de coordenadas, portanto podem ser
agrupadas em uma única realização. Sendo as realizações dos anos 70 estatisticamente iguais.

2. Astro Datum Chuá (1961)
- Época de Referência: 1961
- Elipsoide: Internacional 1924
	- Semieixo maior: 6.378.388 m
	- Achatamento: 1/297
- Origem: Ajustado a observações astronômicas na região de Chuá (MG) -> Topocêntrico
- Observação: Teve uso limitado e regional, principalmente em Minas Gerais. NÃO FOI OFICIALMENTE USADO, FOI TIDO COMO UM ENSAIO PARA O SAD69.

3. SAD69 (South American Datum 1969) : Foi adotado oficialmente no Brasil em 1977 em susbtituição ao Córrego Alegre.
- Época de Referência: 1969
- Elipsoide: Internacional 1967
	- Semieixo maior: a = 6.378.160 m (menor do que CA)
	- Achatamento: f = 1/298,25 (menor que CA)
- Origem: Posicionado astronômico no vértice de Chuá (MG). Ajustado ao continente sul-americano (topocêntrico)
- Observação: Foi o sistema oficial do Brasil de 1979 até 2015, amplamente usado em cartografia.
- Ondulação geoidal: N = 0
- Coordenadas do ponto de origem:

#	"A(latitude astronômica) = 19"45'41.34"S e B(longitude astronômica) = 48"06'07.80"W
#	"A(astro-geodésico) = 19"45'41.6527"S  e   B(longitude astro-geodésica) = 48"06'04.0639"W  ( desvio da vertical )
#"	AZA = 271"30'05.42" e AZ = 271"30'04.05

Os componentes de desvio da vertical do datum SAD69 foram determinados por métodos astro-geodésicos como:
componente meridiana ξ= 0,31" e η (componente vertical)η = -3,52" de arco. 

Sendo assim, o ponto da geoide com a elipsoide não são paralelos em função do desvio da vertical.

As realizações foram ocorridas na décadas de 70, cobrindo grande parte do território brasileiro com dados GPS e TRANSIT. Proporcionando um custo menor
e rapidez nos levantamentos. Oferecendo uma melhor qualidade geométrica com novas tecnologias sendo realizado o SAD69/96.

ξΨψωη

ζ= 0,31 ξΨψωη

4. WGS84 (World Geodetic System 1984)
- Época de Referência: 1984 (atualizações periódicas)
- Elipsoide: WGS84
- Semieixo maior: 6.378.137 m
- Achatamento: 1/298,257223563
- Origem: Centro de Massa da Terra (geocêntrico)
- Observação: Sistema global usado pelo GPS; compatível com o SIRGAS2000.

5. SIRGAS2000 (Sistema de Referência Geocêntrico para as Américas, adotado em 2015): Foi adotado com a motivação de um referencial único para o 
continente Sul-americano com técnicas modernas com 67 estações distribuídas em 11 países da América do Sul. Posteriormente, o projeto SIRGAS foi
expandido para a América Central, do Norte e Caribe. A segunda realização da rede continental ocorreu em maio de 2000 contando com 184 estações
estabelecidas em países das três Américas, sendo 21 delas localizadas no Brasil. Atualmente, o SIRGAS está materializado por uma rede de estações
GNSS de monitoramente contínuo (SIRGAS-CON) formada por cerca de 400 estações.
- Época de Referência: 2000,4
- Elipsoide: GRS80
- Semieixo maior: a = 6.378.137 m
- Achatamento: f = 1/298,257222101
- Origem: Centro de Massa da Terra (Geocêntrico)
- Observação: Sistema oficial do Brasil desde 25/02/2015; compatível com WGS84.
- Estações de referência: redes planimétricas em 21 etações da rede continental SIRGAS2000.

Assim, a evolução foi: Córrego Alegre → Astro Datum Chuá → SAD69 → WGS84 → SIRGAS2000."""
    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("GEOPROCESSAMENTO", level=1)

    p1 = (
    """
    Geoprocessamento são todas as tecnologias utilizadas para aquisição, processamento, armazenamento, manutenção, interpretação e/ou análise de
dados e informações georreferenciadas (DOMINGUES, 2007) 
Um dado espacial descreve um fenômeno associado a alguma dimensão espacial. Um dado geográfico ou geoespacial é um dado espacial em que a
dimensão espacial está associada à sua localização na superfície terrestre, em determinado instante ou período de tempo.
Entre essas tecnologias e métodos se destacam: topografia, cartografia digital, SIG (Sistema de Informação Geográfica), CAD (Computer Aided Design),
GNSS (Global Navigation Satelite System), Sensoriamento Remoto de imagens orbitais (satélite) e não orbitais (fotogrametria)
    
    O SIG (Sistemas de informação geográfica) é a tecnologia que une a informação geográfica à informação descritiva. 
O SIG possibilita a capacidade de visualizar, manipular, analisar e transformar informações geográficas. 
Os componentes de um SIG existentes assim como as máquinas(hardware), os softwares (algortimos, scripts), banco de dados (dataware), pessoas qualificadas, 
dados (localização dos rios, territórios, matas), métodos e procedimentos em saber aplicar e transformar os dados em valor são o que definem o termo GEOPROCESSAMENTO.

    
    Os principais tipos de dados geográficos são:
Dados de referência e cadastrais: é a parte espacial de referência para o SIG, armazenada em forma de coordenadas, podendo ser vetorial ou matricial,
e seus atributos não gráficos são armazenados em um banco de dados. 

Dados temáticos: admitem tanto representação matricial quanto vetorial, e são dados referentes à temática a ser abordada no SIG, podendo ser dados:
estatísticos, de vegetação, de uso do solo, de geologia, entre outros.

Redes: são parte dos dados de referência e temáticos, que são armazenados em forma de coordenadas vetoriais, com a topologia arco-nó e
seus atributos não gráficos são guardados em um banco de dados.

Imagens de sensoriamento remoto: são insumos tanto para mapeamento de referência, quanto para mapeamento temático, e são armazenadas em
representação matricial.
Modelos numéricos de terreno: são gerados por meio de algoritmos e podem ser armazenados em grades regulares (representação matricial), grades
triangulares (representação vetorial com topologia arco-nó) ou isolinhas (representação vetorial com topologia). 

As principais funções de um SIG são:
1. Banco de dados geográficos / 2. Gerencia de dados espaciais / 3. Consulta e análise/visualização(plotagem)/Entrada/integração de dados/ 4. Interface
    """
    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p2 = (
    """
    A estrutura do SIG é composta por CAMADAS. Todos os sotfwares, hoje em dia, são compostas por Layers. Cada camada, cada elemento da paisagem
representa várias informação. Cada camada, cada layers poderá ser alterado. Uma carta topográfica, por exemplo, todas as informações estão mapeadas em uma só
camada. No SIG, podemos ir além, aprofundar em cada camada. Existem 2 tipos básicos de camadas (layers), um vetor e o raster.
	Um arquivo vetorial são todos os dados representados por pontos, linhas e áreas (polígono), ainda possível, colocar atributos nos vetores, em forma de tabelas.
Um ponto vetorial pode ser uma árvore, uma casa. Um arquivo vetorial de linha poderia ser um rio, um córrego. Um vetor tipo área poderia ser um lago, um sítio.
Em resumo, dados vetoriais representam geometrias que são pontos(marcadores de locais no espaço [par de coordenadas associado][sem dimensões, sem área], linhas e polígonos(áreas)
As linhas e polígonos possuem espaço. As linhas são unidimensionais (comprimento) e permitem também análises de ocupação no espaço.

	Os formatos de arquivos de layers vetoriais são os: shapefile(.shp) o mais universal. Podendo ser por tabelas (.tab) e o formatos CAD(.dwg;.dgn;.dfx)
Mas o específico para SIG é o SHAPEFILE.(.shp) O qual é formado por três arquivos, o ".shp", que indica a geometria, o '.dbf', que contém a tabela de atributos e o '.SHX', indexador de arquivo. ( que vai indexar o 'shp' ao 'dbf' )
Tendo esses três arquivos, qualquer softwares de processamento irá reconhecer. Salvando o arquivo completo, gera-se o arquivo '.prj'.

#imagem # imagem

Elementos do arquivo '.dbf':

	É uma tabela com colunas e as linhas com seus atributos. Um dos elementos serão uma coluna com os tipos de geometria e o valor da geometria.
Em termos de prova, o Google Earth não é um SIG. Os tipos de dados, além da geometria, podem vim: 'string' -> Dados textuais / 'Integer' -> Números inteiros
/ 'float/double' -> Números Decimais / 'byte' -> número inteiro de 6 bits (0 a 255) 
	Os arquivos de RASTER é outro formato em que são imagens geradas por satélites, radares, foto aéreas. Só que para se tornar um RASTER, a imagem
deverá ser GEORREFERENCIADA a um sistema de coordenadas. Uma referência espacial. Sem essas informações, ela será uma simples imagem.
	Os formatos são: GeoTIFF(.tif) <- Geradas imagens por satélites // JPEG(.jpg) // bitmap(.bmp). O mais conhecido mesmo é o TIFF.
	Os dados a serem funcionais precisam ter uma referência espacial. Ou seja, são layers associados a sistemas de coordenadas e projeções cartográficas.
Justamente para medir distâncias, definir escalas, medir áreas. Sem referência espacial não há possibilidade de fazer análises no SIG.
Os valores numéricos dos arquivos raster são atribuidos aos Pixels em que permite a realização de operações aritméticas.
Os arquivos de dados 'raster' é composta por uma matriz contínua de dados com valores numéricos atribuídos aos Pixels.

    """
    )
    doc.add_paragraph(p2).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p3=(
    """
    Arquivos KML/KMZ:
Um SIG também aceita arquivos KML o qual estão mais associados ao google earth , ( KML -> keyhole Markup Language ), o KMZ agrupa varios kmls dentro dele.
O formato é open source e agrega atributos de elementos, traz uma estrutura de metadados mas não funciona como tabela de banco de dados. (.db)
Interoperabilidade do arquivo permite usar em sofwares de SIG, ele é nativo do Google Earth. Mesmo podendo ser gerado por SIGs também.

    Arquivos geopackage:

	São arquivos de código aberto (OPEN GEOSPATIAL CONSORTIUM) em que um único arquivo pode armazenar diversas camadas de dados espaciais e atributos.
Baseado em SQlite. Inclusive, 'raster'.
As coordenadas em que falamos são as coordenadas geográficas ( latitudes e longitudes (GMS) ) e as coordenadas planas ou cartesianas (UTM - Universal Transversa de Marcator )
"""
    )
    doc.add_paragraph(p3).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p4= ("""
       

Em resumo, as vantagens do SIG é centralizar as funções de gestão de banco de dados, realizar relações topológicas e espaciais, relações matemáticas (relação de diferentes atributos, elementos geográficos) e geoestatística (correlação espacial)
	Com o SIG podemos gerar mapas digitais e físicos. Medir as distâncias e áreas entre pontos, entre linhas e permite calcular áreas.
	Gerar bases cartográficas vetoriais como pontos ( cidades, municípios, parada de ônibus ), linhas ( rodovias, falhas geológicas, rios, córregos )
e áreas ( rios, lagos, florestas, lotes, limites de cidades e estados )
Assim, os dados poderão ser transformados em novas infomações. O geoprocessamento de dados vetoriais, por exemplo,
recorte, união, raios, extração de áreas. ( apagar um limite sobreposto a outro )
Os dados matriciais, os rasters, advindos de sensores remotos possibilitam a fazer álgebra de mapas e reclassificação. 
O SIG permite também fazer análises 3D. Visualizações trimensionais (MDE - Modelos digitais de elevação).
Cálculos de volume ( cruzamento de área ) / Calcular as distâncias topográficas ( leva em consideração as deformidades do relevo )
Gerenciamento de dados geográficos ( Consultas espaciais ( atributos dos dados -> Descrição dos dados geográficos ); consultas de localização entre objetos de regiões. 
Geração de simbologias ( Atribuir legendas, convenções Ex: prédios acima de 7 pavimentos, bairros acima de 1000 habitantes )

	Os softwares mais conhecidos:

	Os softwares proprietários são pagos: ArcGIS e MaphInfo. (ESRI)
	Os softwares gratuito: SPRING  (INPE) [ voltado para o sensoriamento remoto ]
	Os softwares livres: alterar códigos e colaborativo (plugins [python]( QGIS e gvSIG ) <- EMBRAPA, INCRA (FOSS)
    """
    # imagem
    )
    doc.add_paragraph(p4).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Conclusão
    texto_conclusao = (
        "Os Sistemas Gerenciadores de Banco de Dados são indispensáveis para o funcionamento de praticamente todas as "
        "aplicações modernas. Eles asseguram integridade, segurança e desempenho no gerenciamento das informações, "
        "permitindo que organizações tomem decisões mais assertivas e baseadas em dados concretos. A evolução constante "
        "dessas tecnologias, especialmente com o avanço da computação em nuvem e da inteligência artificial, aponta para "
        "um futuro em que os SGBDs serão ainda mais automatizados, escaláveis e integrados a diferentes plataformas."
    )
    adiciona_secao_titulo_e_paragrafo(doc, "Conclusão", texto_conclusao)

    # Referências
    doc.add_heading("Referências", level=1)
    referencias = [
        "DATE, C. J. Introdução a Sistemas de Bancos de Dados. 8ª edição. Rio de Janeiro: Elsevier, 2004.",
        "ELMASRI, R.; NAVATHE, S. B. Sistemas de Banco de Dados. 7ª edição. Pearson, 2019.",
        "SILBERSCHATZ, A.; KORTH, H. F.; SUDARSHAN, S. Sistemas de Banco de Dados. 6ª edição. AMGH, 2020.",
        "Oracle Corporation. Oracle Database Documentation. Disponível em: https://docs.oracle.com",
        "PostgreSQL Global Development Group. PostgreSQL Documentation. Disponível em: https://www.postgresql.org/docs/"
    ]
    for ref in referencias:
        # cada referência em parágrafo separado com marcador de hífen no início
        p = doc.add_paragraph(f"- {ref}")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Salvar arquivo
    doc.save(nome_arquivo)
    os.startfile(nome_arquivo)
    print(f"Documento gerado com sucesso: {nome_arquivo}")
    adiciona_sumario(doc,main)


    
      
if __name__ == "__main__":
    main()

"""
gera_sgbd_docx.py
Gera um documento Word com o trabalho:
"Um levantamento técnico sobre Sistemas Gerenciadores de Banco de Dados"

Como usar:
1. Instale a biblioteca python-docx:
   pip install python-docx

2. Execute:
   python gera_sgbd_docx.py

O arquivo será salvo no mesmo diretório com o nome:
Um_levantamento_tecnico_sobre_SGBD_Daniel_Martins_Franca.docx
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

def cria_capa(doc, titulo, autor, curso, instituicao, local_data):
    # inserir espaçamento superior na capa
    p_top = doc.add_paragraph()
    p_top.add_run("\n\n\n\n\n\n\n")
    # título centralizado e em negrito
    p_titulo = doc.add_paragraph()
    run = p_titulo.add_run(titulo)
    run.bold = True
    run.font.size = Pt(18)
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # espaçamento
    doc.add_paragraph("\n\n\n\n")

    p_autor = doc.add_paragraph(autor)
    p_autor.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_curso = doc.add_paragraph(curso)
    p_curso.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_instituicao = doc.add_paragraph(instituicao)
    p_instituicao.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_local = doc.add_paragraph(local_data)
    p_local.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

def aplica_fonte_padrao(doc, nome_fonte="Times New Roman", tamanho=12):
    """
    Define a fonte padrão do documento (para parágrafos novos).
    Nota: python-docx não altera facilmente estilos existentes em todos os casos,
    mas podemos ajustar os estilos principais.
    """
    style = doc.styles['Normal']
    style.font.name = nome_fonte
    # necessário para compatibilidade com alguns sistemas:
    r = style._element.rPr.rFonts
    r.set(qn('w:eastAsia'), nome_fonte)
    style.font.size = Pt(tamanho)

def adiciona_secao_titulo_e_paragrafo(doc, titulo, texto):
    doc.add_heading(titulo, level=1)
    parag = doc.add_paragraph(texto)
    parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def main():
    titulo = "Noções básicas de Geodésia"
    autor = "Daniel Martins França"
    curso = "Curso de Extensão"
    instituicao = "Faculdade Metropolitana"
    local_data = "Brasília – DF, novembro de 2025"
    nome_arquivo = "geografia.docx"

    doc = Document()
    aplica_fonte_padrao(doc, nome_fonte="Times New Roman", tamanho=12)

    # Capa
    cria_capa(doc, titulo, autor, curso, instituicao, local_data)

    # Introdução
    texto_intro = (
    "Diversos modelos foram adotados ao longo da história. Não existe modelo errado, tudo depende do contexto histórico e da aplicação."
	"Assim, o primeiro modelo é chamado de geoidal, é o mais aproximado da forma real, podendo ser determinado pelas medidas gravimétricas," 
"ou seja, medidas da força da gravidade, explicado a seguir. O geoide não pode ser definido matematicamente pois é afetado pelas variações da densidade dos"
"elementos constituintes da crosta terrestre. Além da distribuição irregular das massas terrestres e oceânicas."
"	Devido às irregularidades da superfície terrestre, utilizam-se modelos para a sua representação, mais simples, outros regulares e geométricos e que se"
"aproximam da forma real para efetuar os cálculos. Cada modelo tem a sua aplicação e quanto mais complexa a figura empregada para a representação da Terra,"
"mais complexo serão os cálculos sobre esta superfície. A forma da Terra gira em torno do seu eixo e movendo-se dentro do sistema solar. E o resultado da interação"
"de forças internas e externas tais como: gravidade, força centrífuga e a constituição diferente dos materiais que a formam a geoide ao longo de milhares de anos."
"As forças tectônicas provocam modificações na superfície, que se traduzem por irregularidades topográficas, sobre as quais são realizados os mapeamentos, "
"medições e estudos das mais variedades."
"."
    )
    adiciona_secao_titulo_e_paragrafo(doc, "Estudos Geodésicos:", texto_intro)

    # Desenvolvimento (vários parágrafos)
    doc.add_heading("Geoide:", level=1)

    p1 = (
        "É a forma física real, que sofre frequentes alterações devido a natureza (campo gravitacional do planeta terra, movimentos tectônicos, condições climáticas,"
"erosões, etc) e à ação do homem, portanto, não serve para definir forma sistemática da Terra. Portanto, a zona de contato da superfície terrestre topográfica e"
"o geoide que defini-se o nível zero das altitudes. Em resumo, a força gravitacional que age sobre a terra irá definir a forma irregular do planeta. O equilíbrio"
"desse potencial gravitacional é o que gera a forma física da Terra. Mantendo os oceanos e a divisão dos continentes com uma relação de equilíbrio do planeta."
"	Portanto, a geoide é a superfície equipotencial gravitacional que mais se aproxima da superfície formada pelo prolongamento dos oceanos (nível médio dos mares)"
"sob os continentes. Essa supefície sofre variações conforme ocorrem alterações no campo gravitacional terrestre e, portanto, não segue leis matemáticas que permitam"
"um modelamento da Terra. Ainda assim é empregado como referência para a determinação das altitudes. "

    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p2 = (
        "A Gravimetria é um método de medida da gravidade utilizado em diversos pontos distribuidos do planeta Terra. Para que apartir dessas medições se alcançem um modelo"
"de ação gravitacional equilibrado. Portanto, a Gravimetria é um método da Geodésia física que determina os níveis do campo gravitacional da Terra e, com isto,"
"determinar o geoide. A densidade de pontos é muito importante para a determinação do geoide. Quanto mais pontos de medição gravimétrica existirem na superfície" 
"terrestre, mais precisa é a figura geoidal. A incidência nos pontos gravitacionais é perpendicular a superfície da terra deverá ser ortométrica (distância contada sobre a vertical)"
"As altitudes ortométricas (ângulo reto) a ação da gravidade nesses pontos em que teremos a determinação de uma superficíe equipontencial. Essa superfície"
"equipotencial é chamada de um modelo geoide. Em média, coincide com o valor médio do nível médio das águas do mar, por isso é usado para medições"
"de altitudes  (altimetria). A superfície geoidal é mais irregular que qualquer outra superfície. A superfície varia entre os +8.850 m (Monte Everest) e -11.000m "
"(Fossa das Marianas). O geoide varia apenas cerca de +-100 m além das superfície do elipsoide de referência."

    )
    doc.add_paragraph(p2).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("Elipsoidal:", level=1)

    p1 = (
        "Como o geoide é uma superfície matematicamente indefinida, as reduções ou transferências de dados a ele são inconsistentes, e para um mapeamento preciso de "
"grandes áreas é necessária a consideração de uma figura geométrica regular. Assim, chega-se na figura matematicamente definida como um elipsoide."
"A elipsoide é a representação matemática mais próximas da forma da Terra. Esse método é utilizado para GNSS e Geodésia Celeste, inclusive. (tecnologias novas)"
"O modelo Elipsoide de revolução é o nome que se dá a uma figura gerada pela rotação de uma elipse sobre um de seus eixos(de norte a sul da esfera). Eixo polar (Norte e Sul) "
"é menor do que o eixo equatorial ( horizontal ) que por sua vez também possui semi-eixo. A relação entre o semi-eixo maior (equatorial) e o semi-eixo menor (polar)"
"é o que define a chamada excêntricidade em que ocorre o achatamento da elipsoide. A posição deste elipsoide em relação à Terra, bem como sua forma e tamanho,"
"constituem um conjunto de parâmetros definidos a partir do elipsoide ao seu ajustamento com a geoide onde são usualmente denominados Datum Geodésico. "
"Uma elipsoide de revolução pode ser determinado para melhor se adaptar a uma região, país ou continente, evitando a ocorrência de desníveis geoidais muito exagerados."
"Em geral, cada país adotou um elipsoide como referência para os trabalhos geodésicos e topográficos. Atualmente existem pelo menos 50 elipsoides existentes"
"utilizados pelo mundo, sendo alguns mais conhecidos: SAD-69, WGS-84, SIRGA-2000. "
"Essas siglas, sempre encontradas no mapas, são elipsoides de revolução. Elas ilustram a referência utilizada para usar em trabalhos cartográficos."
"Sendo assim, a superfície elipsoidal, que é obtido a partir do elipsoide de revolução, é uma figura matemática constante. Então, cada país ou região, para estudos"
"de mapeamento, por exemplo, é estabelecer qual é o elipsoide de referência para o mapeamento naquele local."
"	O semi-eixo do WGS 1984 possui um eixo equatorial, em metros, 6.378.137. O eixo polar possui, em metros, 6.356.752.3142."
"Foram usados equipamentos novos como GPS, GNSS, RTK para o datum WGS 1984."

    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("DATUM:", level=1)
    p1 = (
        "A posição do elipsoide em relação a terra, bem como sua forma e tamanho, constituem um conjunto de parâmetros que são usualmente denominados de DATUM"
"GEODÉSICOS. Datum corresponde a um ponto ou plano de referência para levantamento verticais e horizontais, os quais estabelecem as posições de feições sobre"
"a terra."
    )
    doc.add_paragraph(p1).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p5 = (
        "Além disso, os SGBDs possuem mecanismos robustos de controle de acesso e segurança, garantindo que apenas "
        "usuários autorizados possam manipular determinadas informações. Outro ponto fundamental é a capacidade de "
        "recuperação em caso de falhas, através de backups automáticos e logs de transações que permitem restaurar o "
        "banco de dados ao seu estado anterior a um problema."
    )
    doc.add_paragraph(p5).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Conclusão
    texto_conclusao = (
        "Os Sistemas Gerenciadores de Banco de Dados são indispensáveis para o funcionamento de praticamente todas as "
        "aplicações modernas. Eles asseguram integridade, segurança e desempenho no gerenciamento das informações, "
        "permitindo que organizações tomem decisões mais assertivas e baseadas em dados concretos. A evolução constante "
        "dessas tecnologias, especialmente com o avanço da computação em nuvem e da inteligência artificial, aponta para "
        "um futuro em que os SGBDs serão ainda mais automatizados, escaláveis e integrados a diferentes plataformas."
    )
    adiciona_secao_titulo_e_paragrafo(doc, "Conclusão", texto_conclusao)

    # Referências
    doc.add_heading("Referências", level=1)
    referencias = [
        "DATE, C. J. Introdução a Sistemas de Bancos de Dados. 8ª edição. Rio de Janeiro: Elsevier, 2004.",
        "ELMASRI, R.; NAVATHE, S. B. Sistemas de Banco de Dados. 7ª edição. Pearson, 2019.",
        "SILBERSCHATZ, A.; KORTH, H. F.; SUDARSHAN, S. Sistemas de Banco de Dados. 6ª edição. AMGH, 2020.",
        "Oracle Corporation. Oracle Database Documentation. Disponível em: https://docs.oracle.com",
        "PostgreSQL Global Development Group. PostgreSQL Documentation. Disponível em: https://www.postgresql.org/docs/"
    ]
    for ref in referencias:
        # cada referência em parágrafo separado com marcador de hífen no início
        p = doc.add_paragraph(f"- {ref}")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Salvar arquivo
    doc.save(nome_arquivo)
    os.startfile(nome_arquivo)
    print(f"Documento gerado com sucesso: {nome_arquivo}")

if __name__ == "__main__":
    main()
